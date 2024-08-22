import streamlit as st

import requests
import pandas as pd
from io import StringIO, BytesIO
import base64
import os
import ast
import re
from docx import Document
from dotenv import load_dotenv
import markdown2
import markdown2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from docx.shared import Pt


class init:
    def main():
        load_dotenv()
    
    def markdown_para_docx(markdown_texto, caminho_saida):
        # Converte o markdown para HTML
        html_texto = markdown2.markdown(markdown_texto)
        
        # Cria um novo documento
        documento = Document()
        
        # Adiciona conteúdo ao documento com formatação
        adicionar_html_ao_documento(documento, html_texto)
        
        # Salva o documento no caminho especificado
        documento.save(caminho_saida)

def adicionar_html_ao_documento(documento, html_texto):
    soup = BeautifulSoup(html_texto, 'html.parser')

    for elemento in soup:
        if elemento.name == 'p':
            paragrafo = documento.add_paragraph()
            adicionar_html_ao_paragrafo(paragrafo, elemento)
        elif elemento.name == 'h1':
            paragrafo = documento.add_paragraph()
            run = paragrafo.add_run(elemento.text)
            run.bold = True
            run.font.size = Pt(24)
        elif elemento.name == 'h2':
            paragrafo = documento.add_paragraph()
            run = paragrafo.add_run(elemento.text)
            run.bold = True
            run.font.size = Pt(20)
        # Adicione mais elifs para outras tags HTML, se necessário

def adicionar_html_ao_paragrafo(paragrafo, elemento):
    for subelemento in elemento:
        if subelemento.name == 'strong':
            run = paragrafo.add_run(subelemento.text)
            run.bold = True
        elif subelemento.name == 'em':
            run = paragrafo.add_run(subelemento.text)
            run.italic = True
        else:
            run = paragrafo.add_run(subelemento if isinstance(subelemento, str) else subelemento.text)

class debugger:
        def console_debugging_status(debug):
            print()
            print(debug)
            print('-'*30)
class api:
    def inia(url, archive):
                
        # Abrir o arquivo PDF no modo leitura binária
                files = {'file': (archive.name, archive.getvalue(), archive.type)}
                    # Fazer a solicitação POST para o webhook
                response = requests.post(url, files=files)
                    # Verificar se a resposta indica sucesso 
                debugger.console_debugging_status('Arquivo Enviado com sucesso')

                return response.text

    def main(url, archive):
        
                    response_webhook = api.inia(url, archive=archive).split('[SPLIT]')
                    planilhas,relatorio,dados_paciente = response_webhook[0], response_webhook[1], response_webhook[2]
                    planilhas = planilhas.replace("```","").replace('```','')
                    planilhas = StringIO(planilhas)
                    df = pd.read_csv(planilhas,delimiter = ';')
                    debugger.console_debugging_status('Planilha Gerada com sucesso')

                    #Dict com metadados
                    #dados_paciente = dados_paciente.replace('```json',"").replace("```","")
                    #dados_paciente_dict = ast.literal_eval(dados_paciente)
                    #console_debugging_status('Metatados Extraidos com Sucesso')

                    #Separa o CSV entre anormais e normais
                    df_anormal = df[~df.iloc[:, 4].str.contains('acima|abaixo') == False]
                    df_normal = df[~df.iloc[:, 4].str.contains('acima|abaixo')]
                    debugger.console_debugging_status('Planilhas divididas com sucesso')
                    
                
                    #Gera arquivo xlsx a partir dos csv's
                    normal_output = BytesIO()
                    anormal_output = BytesIO()
                    diagnostic_output = BytesIO()
                    df_normal.to_excel(normal_output, index=False)
                    df_anormal.to_excel(anormal_output, index=False)
                    init.markdown_para_docx(relatorio, diagnostic_output)

                    st.download_button("Download Normal File", data=normal_output.getvalue(), file_name='normal_file.xlsx') 
                    st.download_button("Download Anormal File", data=anormal_output.getvalue(), file_name='anormal_file.xlsx')
                    st.download_button("Download Relatorio ", data= diagnostic_output.getvalue(),file_name='relatorio.docx')
                    print('Concluido')


class webapp:
        def main(url):  
            st.set_page_config("Inia Desktop")
            pdf = st.file_uploader("Faça seu upload aqui", type=['pdf'])
            if pdf:
                api.main(url, pdf)
#st.download_button("Baixe o PDF")
init.main()
webapp.main(os.environ.get('INIA_API'))
#class api, webapp
#def api.inia, api.main, webapp.main, webapp.data_manager